# -*- coding: utf-8 -*-
"""
parsers.py — Implementación del Parser (PDF / DOCX / IMG → IR)

Resumen
-------
Convierte documentos heterogéneos a una representación intermedia (IR) homogénea
basada en JSON/MD, con bloques de texto y tablas por página, manteniendo un
contrato de salida estable para los siguientes subsistemas (chunking, IE, etc.).

Este módulo ha sido MEJORADO para:
- Reconstruir párrafos lógicos en PDF (evitar “micro-oraciones” por salto visual).
- Aplicar heurísticas duales: textual (puntuación/capitalización) y layout-aware (vertical gaps).
- Conservar trazabilidad fina: `source_lines`, `prov` con extractor compuesto y notas.
- Enriquecer metadatos y métricas operables: `layout_loss`, `fusion_rate`, `n_blocks_raw`, `n_blocks_final`.
- Robustecer detección de headings y list items.
- Mantener backward-compat con contratos existentes (`schemas.py`).

Características clave
---------------------
- Detección de tipo por MIME / extensión y dispatch a parser especializado.
- PDF: reconstrucción de párrafos (texto + gaps verticales) y tablas básicas con pdfplumber.
- DOCX: párrafos/headings + tablas con python-docx.
- IMG: OCR con pytesseract (opcional).
- Fallback OCR para páginas PDF sin texto/tabla (típico en PDFs escaneados).
- Normalización de texto configurable (espacios, dehyphenate).
- Metadatos útiles (size_bytes, sha256, page_count) y provenance para trazabilidad.
- Heurísticas opcionales:
  * list_item: detecta bullets/guiones y etiqueta como 'list_item'
  * heading: marca encabezados simples cuando no hay estilos (PDF)

Parámetros importantes
----------------------
- ocr_lang: Idiomas OCR para Tesseract (ej. "spa", "eng", "spa+eng").
- ocr_resolution: DPI para rasterizar páginas PDF en fallback OCR.
- normalize_whitespace: Colapsa espacios múltiples y limpia líneas.
- dehyphenate: Une palabras cortadas por guion ("infor-\\n mación" -> "información").
- enable_pdf_ocr_fallback: Activa OCR por página si pdfplumber no extrajo texto/tabla.
- enable_pdf_heading_heuristics: Heurística conservadora para headings en PDF.
- enable_list_item_detection: Detecta bullets/guiones como list items.
- enable_lang_detect: Detecta idioma (doc/página) si está disponible `langdetect`.
- tesseract_cmd: Ruta al ejecutable de tesseract (útil en Windows).

Mejoras futuras (ideas)
-----------------------
- Headings PDF robustos con font-size/weight vía page.extract_words() (fontname/size).
- Tablas PDF robustas con camelot o tabula-py (cuando el PDF es vectorial).
- Preproceso de imagen (deskew/denoise/binarización) antes de OCR (OpenCV).
- Segmentación por columnas para PDFs multi-columna (detectar clusters X).
"""

from __future__ import annotations
import os
import io
import mimetypes
import logging
import hashlib
import re
from statistics import median
from typing import List, Dict, Any, Optional, Tuple

import pdfplumber

# Import opcionales protegidos: no romper si no están presentes
try:
    import docx  # python-docx para .docx/.doc
except ImportError:
    docx = None

try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract, Image = None, None

# Detección de idioma opcional
try:
    from langdetect import detect as _langdetect
except Exception:
    _langdetect = None

from parser.schemas import (
    DocumentIR, PageIR, TextBlock, TableBlock, TableCell,
    Provenance, OCRInfo
)

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# ---------------------------------------------------------------------
# Utilidades internas
# ---------------------------------------------------------------------

# Bullets/viñetas comunes para detectar list items
_BULLETS = tuple("•·◦▪-–—*·")

# Fin de oración aproximado (resistente a comillas/ellipses)
_SENT_END = re.compile(r'[\.!?…]"?$')

def _is_list_item_text(s: str) -> bool:
    """Heurística simple: detecta líneas con bullets/guiones al inicio (tras espacios)."""
    if not s:
        return False
    ls = s.lstrip()
    return bool(ls and ls[0] in _BULLETS)

def _guess_mime(path: str) -> str:
    """
    Infiera el MIME a partir de la extensión.
    Si no se conoce, usa 'application/octet-stream' como fallback.
    """
    mime, _ = mimetypes.guess_type(path)
    return mime or "application/octet-stream"

def _sha256(path: str, chunk_size: int = 1024 * 1024) -> str:
    """Calcula sha256 en streaming para trazabilidad."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        while True:
            chunk = f.read(chunk_size)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()

def _detect_lang(text: str) -> str:
    """
    Detecta idioma 'es'/'en' si langdetect está disponible, si no 'und'.
    Nota: puedes extender mapeos a pt/fr/de si lo necesitas.
    """
    if not text or _langdetect is None:
        return "und"
    try:
        code = _langdetect(text)
        return code if code in {"es", "en", "pt", "fr", "de"} else "und"
    except Exception:
        return "und"

# ---------------------- Reconstrucción de párrafos ----------------------

def _normalize_lines_for_merge(lines: List[str], dehyphenate: bool, normalize_whitespace: bool) -> List[str]:
    """
    Normaliza líneas previas a la fusión de párrafos (igual política que _normalize_text).
    No colapsamos los saltos aquí; solo limpiamos cada línea.
    """
    norm = []
    for s in lines:
        if not isinstance(s, str):
            norm.append(s)
            continue
        s = s.replace("\xa0", " ").replace("\r", "")
        if dehyphenate:
            s = s.replace("-\n", "").replace("-\r\n", "")
        if normalize_whitespace:
            s = " ".join(s.split())
        norm.append(s.strip())
    return norm

def _merge_lines_textual(lines: List[str]) -> Tuple[List[str], List[List[int]]]:
    """
    Reconstruye párrafos lógicos con heurística TEXTUAL pura.
    Devuelve:
      - lista de párrafos (strings)
      - lista de índices de líneas (source_lines) que componen cada párrafo
    Reglas:
      - Se respeta doble salto como separador de párrafo (líneas vacías).
      - Si la línea termina en fin de oración (_SENT_END) y la siguiente inicia con mayúscula o está vacía → corte.
      - Si no hay siguiente línea, se corta.
      - Durante la unión se inyecta espacio simple.
    """
    paras: List[str] = []
    sources: List[List[int]] = []
    buf: List[str] = []
    src: List[int] = []

    n = len(lines)
    for i in range(n):
        l = lines[i].strip()
        if not l:
            if buf:
                paras.append(" ".join(buf).strip())
                sources.append(src[:])
                buf.clear()
                src.clear()
            continue

        buf.append(l)
        src.append(i)

        next_line = lines[i + 1].strip() if i + 1 < n else ""
        end_here = bool(_SENT_END.search(l))
        next_is_capital = bool(next_line and next_line[0].isupper())

        # Cortes “naturales”: fin de oración + próxima capitalizada o no hay siguiente
        if end_here and (not next_line or next_is_capital):
            paras.append(" ".join(buf).strip())
            sources.append(src[:])
            buf.clear()
            src.clear()

    if buf:
        paras.append(" ".join(buf).strip())
        sources.append(src[:])

    return paras, sources

def _merge_lines_layout_aware(page: pdfplumber.page.Page, lines: List[str]) -> Tuple[List[str], List[List[int]]]:
    """
    Intenta reconstruir párrafos usando también “gaps” verticales entre words.
    Estrategia:
      - Usa page.extract_words() para obtener y-centroid/ascender-descender y orden natural.
      - Calcula distancia vertical (delta_y) entre líneas consecutivas “sintetizadas”.
      - Si delta_y > umbral (p.ej. > 1.8 * mediana de gaps) → es separador de párrafo.
      - Si no, une como mismo párrafo.
    Fallback:
      - Si extract_words falla o devuelve poco, se usa la estrategia textual pura.
    """
    try:
        words = page.extract_words(x_tolerance=1, y_tolerance=1, keep_blank_chars=False) or []
    except Exception:
        words = []

    if not words:
        return _merge_lines_textual(lines)

    # Agrupar palabras por “línea” aproximada (mismo y0..y1)
    # pdfplumber ya ordena, pero igual normalizamos a líneas
    synthesized_lines: List[Tuple[float, List[str], int]] = []  # (y_center, words_in_line, line_idx_ref)
    current: List[str] = []
    current_y: Optional[float] = None
    line_map: List[int] = []  # mapea índice sintetizado -> índice original de lines

    # Creamos un índice aproximado original por conteo: asumimos que split("\n") mantuvo orden
    # y cada línea textual corresponde secuencialmente a un bloque de palabras secuente.
    # No es perfecto, pero suficiente para trazabilidad y gaps.
    line_counter = 0
    last_y = None

    for w in words:
        y_center = (w["top"] + w["bottom"]) / 2.0 if ("top" in w and "bottom" in w) else None
        txt = w.get("text", "").strip()
        if txt == "":
            continue
        if current_y is None:
            current_y = y_center
            current = [txt]
            last_y = y_center
            continue
        # Si el delta_y con respecto a la línea actual es pequeño, pertenece a la misma línea
        if y_center is not None and current_y is not None and abs(y_center - current_y) < 2.5:  # tolerancia conservadora
            current.append(txt)
            last_y = y_center
        else:
            # cerramos línea
            synthesized_lines.append((current_y if current_y is not None else 0.0, current[:], line_counter))
            line_map.append(min(line_counter, len(lines) - 1))
            line_counter += 1
            current_y = y_center
            current = [txt]
            last_y = y_center
    if current:
        synthesized_lines.append((current_y if current_y is not None else 0.0, current[:], line_counter))
        line_map.append(min(line_counter, len(lines) - 1))

    if len(synthesized_lines) <= 1:
        return _merge_lines_textual(lines)

    # Calculamos gaps verticales entre líneas sintetizadas
    y_list = [y for (y, _, _) in synthesized_lines]
    gaps = [abs(y_list[i] - y_list[i - 1]) for i in range(1, len(y_list))]
    med_gap = median(gaps) if gaps else 0.0
    threshold = 1.8 * med_gap if med_gap > 0 else 9999  # si no hay señal, casi nunca cortamos por layout

    paras: List[str] = []
    sources: List[List[int]] = []
    buf: List[str] = []
    src: List[int] = []

    for idx, (y, words_in_line, line_ref) in enumerate(synthesized_lines):
        merged_line = " ".join(words_in_line).strip()
        # Emparejamos con la línea textual normalizada (por best-effort)
        textual_idx = min(line_map[idx], len(lines) - 1)
        textual_line = lines[textual_idx].strip()

        # Preferimos la línea textual normalizada para limpieza; si vacía, usamos merged_line
        final_line = textual_line or merged_line
        if not final_line:
            # Si está vacía, consideramos que puede ser separador
            if buf:
                paras.append(" ".join(buf).strip())
                sources.append(src[:])
                buf.clear()
                src.clear()
            continue

        buf.append(final_line)
        src.append(textual_idx)

        # Decidir corte por layout gap o por textual
        gap_next = abs(y_list[idx + 1] - y) if idx + 1 < len(y_list) else None
        next_textual = lines[textual_idx + 1].strip() if textual_idx + 1 < len(lines) else ""
        end_textual = bool(_SENT_END.search(final_line)) and (not next_textual or next_textual[0].isupper())

        cut = False
        if gap_next is None:
            cut = True
        elif gap_next > threshold:
            cut = True
        elif end_textual:
            cut = True

        if cut:
            paras.append(" ".join(buf).strip())
            sources.append(sorted(set(src)))
            buf.clear()
            src.clear()

    if buf:
        paras.append(" ".join(buf).strip())
        sources.append(sorted(set(src)))

    return paras, sources

# ---------------------------------------------------------------------
# Clase principal
# ---------------------------------------------------------------------

class Parser:
    """
    Fachada del Parser de documentos con opciones configurables.

    Args
    ----
    ocr_lang : str
        Idioma(s) para Tesseract. Ej: "spa", "eng" o "spa+eng".
    ocr_resolution : int
        DPI para rasterizar páginas PDF cuando se usa fallback OCR (200–300 recomendado).
    normalize_whitespace : bool
        Si True, colapsa espacios múltiples, quita extra whitespaces y normaliza líneas.
    dehyphenate : bool
        Si True, une palabras cortadas por guion al final de línea ("infor-\\nmación" -> "información").
    enable_pdf_ocr_fallback : bool
        Si True, intenta OCR por página si pdfplumber no extrajo texto ni tablas (PDFs escaneados).
    enable_pdf_heading_heuristics : bool
        Activa una heurística conservadora para marcar encabezados en PDF.
    enable_list_item_detection : bool
        Si True, intenta clasificar líneas con bullets/guiones como 'list_item'.
    enable_lang_detect : bool
        Si True, intenta detectar idioma del documento y por página (si hay lib).
    tesseract_cmd : Optional[str]
        Ruta al ejecutable de Tesseract (útil en Windows). Ej:
        r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"

    Uso
    ---
        parser = Parser(ocr_lang="spa+eng", ocr_resolution=220)
        doc_ir = parser.parse("archivo.pdf")
    """

    def __init__(
        self,
        ocr_lang: str = "spa",
        ocr_resolution: int = 220,
        normalize_whitespace: bool = True,
        dehyphenate: bool = True,
        enable_pdf_ocr_fallback: bool = True,
        enable_pdf_heading_heuristics: bool = True,
        enable_list_item_detection: bool = True,
        enable_lang_detect: bool = False,
        tesseract_cmd: Optional[str] = None,
    ):
        self.ocr_lang = ocr_lang
        self.ocr_resolution = ocr_resolution
        self.normalize_whitespace = normalize_whitespace
        self.dehyphenate = dehyphenate
        self.enable_pdf_ocr_fallback = enable_pdf_ocr_fallback
        self.enable_pdf_heading_heuristics = enable_pdf_heading_heuristics
        self.enable_list_item_detection = enable_list_item_detection
        self.enable_lang_detect = enable_lang_detect

        # Configurar Tesseract manualmente si pasamos ruta (Windows / entornos custom)
        if tesseract_cmd and pytesseract is not None:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    # ---------------------- Helpers internos ----------------------

    def _normalize_text(self, s: str) -> str:
        """
        Normaliza texto de salida para IR:
        - Reemplaza NBSP (\xa0) por espacio y remueve \r
        - Une palabras cortadas por guion (opcional)
        - Colapsa espacios múltiples y limpia bordes (opcional)
        """
        if not isinstance(s, str):
            return s
        s = s.replace("\xa0", " ").replace("\r", "")
        if self.dehyphenate:
            # Heurística simple: elimina "-\n" para unir palabras cortadas al salto de línea
            s = s.replace("-\n", "").replace("-\r\n", "")
        if self.normalize_whitespace:
            # Colapsa espacios (a diferencia de _normalize_lines_for_merge que se aplica por línea)
            s = " ".join(s.split())
        return s.strip()

    def _maybe_heading(self, txt: str) -> bool:
        """
        Heurística MUY conservadora de heading:
        - Línea relativamente corta
        - No termina en puntuación fuerte (.,:;)
        - Proporción de mayúsculas alta respecto a alfabéticos
        """
        if not self.enable_pdf_heading_heuristics:
            return False
        if not txt:
            return False
        if len(txt) > 80:
            return False
        if txt.endswith((".", ":", ";")):
            return False
        letters = sum(c.isalpha() for c in txt)
        uppers = sum(c.isupper() for c in txt)
        return bool(letters and uppers >= 0.5 * letters)

    # ---------------------- API pública ----------------------

    def parse(self, path: str) -> DocumentIR:
        """
        Detecta el tipo de documento y lo parsea a IR.

        Devuelve
        --------
        DocumentIR
            Estructura con doc_id, meta (size_bytes, page_count, sha256), pages[PageIR], etc.
        """
        if not os.path.exists(path):
            raise FileNotFoundError(path)

        mime = _guess_mime(path)
        logger.info("Parsing start | path=%s mime=%s", path, mime)

        # Metadatos útiles para trazabilidad
        meta: Dict[str, Any] = {"filename": os.path.basename(path)}
        try:
            stat = os.stat(path)
            meta["size_bytes"] = stat.st_size
        except Exception:
            pass
        try:
            meta["sha256"] = _sha256(path)
        except Exception as e:
            logger.debug("No se pudo calcular sha256: %s", e)

        doc = DocumentIR(
            doc_id=DocumentIR.new_id(),
            source_path=path,
            mime=mime,
            meta=meta,
            prov=Provenance(extractor="pdfplumber/python-docx/pytesseract", stage="parser"),
        )

        if mime == "application/pdf" or path.lower().endswith(".pdf"):
            pages = self._parse_pdf(path)
        elif mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ) or path.lower().endswith((".docx", ".doc")):
            pages = self._parse_docx(path)
        elif mime and mime.startswith("image/"):
            pages = self._parse_image(path)
        else:
            raise ValueError(f"Tipo no soportado: {mime} (path={path})")

        # Idioma (opcional)
        if self.enable_lang_detect:
            # Doc-level
            try:
                sample_doc = " ".join(
                    (b.get("text", "") if isinstance(b, dict) else getattr(b, "text", ""))
                    for p in pages for b in p.blocks
                )[:2000]
                doc.lang = _detect_lang(sample_doc) if sample_doc else "und"
            except Exception:
                doc.lang = "und"
            # Page-level
            for p in pages:
                try:
                    sample_pg = " ".join(
                        (b.get("text", "") if isinstance(b, dict) else getattr(b, "text", ""))
                        for b in p.blocks
                    )[:1000]
                    p.lang = _detect_lang(sample_pg) if sample_pg else "und"
                except Exception:
                    p.lang = "und"

        doc.pages = pages
        doc.meta["page_count"] = len(pages)
        logger.info("Parsing done | doc_id=%s pages=%d", doc.doc_id, len(doc.pages))
        return doc

    # ---------------------- Parsers especializados ----------------------

    def _parse_pdf(self, path: str) -> List[PageIR]:
        """
        PDF → PageIR[]:
        - Reconstrucción robusta de párrafos (layout-aware + textual).
        - Tablas básicas con pdfplumber.
        - Fallback OCR por página si no hubo texto ni tablas (PDFs escaneados),
          aplicado con Tesseract si está disponible y habilitado.
        """
        pages_ir: List[PageIR] = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                width, height = page.width, page.height
                page_blocks: List[Dict[str, Any]] = []
                meta_page: Dict[str, Any] = {}

                # 1) Texto crudo por líneas (visual)
                raw_text = page.extract_text(x_tolerance=1, y_tolerance=1) or ""
                raw_lines = raw_text.split("\n")

                # Guardamos cantidad de líneas crudas para métrica
                n_raw_lines = sum(1 for l in raw_lines if l.strip())
                meta_page["n_lines_raw"] = n_raw_lines

                # 2) Normalización de líneas antes de fusionar
                norm_lines = _normalize_lines_for_merge(
                    raw_lines,
                    dehyphenate=self.dehyphenate,
                    normalize_whitespace=self.normalize_whitespace,
                )

                # 3) Intento layout-aware; si no hay señal, cae a textual
                paras, sources = _merge_lines_layout_aware(page, norm_lines)

                # 4) Si por algún motivo quedó vacío (documento patológico), usar textual pura
                if not paras:
                    paras, sources = _merge_lines_textual(norm_lines)

                # Métricas de fusión
                n_final_paras = len([p for p in paras if p.strip()])
                meta_page["n_paragraphs_final"] = n_final_paras
                meta_page["fusion_rate"] = round(n_raw_lines / max(1, n_final_paras), 3) if n_final_paras else None
                # layout_loss simple: proporción de bloques marcados como 'unknown' (aquí 0 si no generamos unknowns)
                meta_page["layout_loss"] = 0.0

                # 5) Construcción de bloques (heading/list_item/paragraph)
                for ptxt, src_idx_list in zip(paras, sources):
                    if not ptxt.strip():
                        continue

                    # list item
                    if self.enable_list_item_detection and _is_list_item_text(ptxt):
                        cleaned = ptxt.lstrip()
                        cleaned = cleaned[1:].lstrip() if cleaned and cleaned[0] in _BULLETS else cleaned
                        blk = TextBlock(
                            type="list_item",
                            text=cleaned,
                            prov=Provenance(extractor="pdfplumber+merge", stage="parser",
                                            notes=f"source_lines={src_idx_list}")
                        )
                    # heading
                    elif self._maybe_heading(ptxt):
                        blk = TextBlock(
                            type="heading",
                            text=ptxt,
                            level=2,
                            prov=Provenance(extractor="pdfplumber+merge", stage="parser",
                                            notes=f"source_lines={src_idx_list}")
                        )
                    # párrafo normal
                    else:
                        blk = TextBlock(
                            type="paragraph",
                            text=ptxt,
                            prov=Provenance(extractor="pdfplumber+merge", stage="parser",
                                            notes=f"source_lines={src_idx_list}")
                        )
                    page_blocks.append(blk.model_dump())

                # 6) Tablas básicas (nota: PDFs escaneados normalmente NO tendrán tablas extraíbles)
                tables = []
                try:
                    tables = page.extract_tables(
                        table_settings={"vertical_strategy": "lines", "horizontal_strategy": "lines"}
                    )
                except Exception as e:
                    logger.debug("Table extraction failed on page %d: %s", i, e)

                for t in tables or []:
                    tb = TableBlock(prov=Provenance(extractor="pdfplumber", stage="parser"))
                    for r_idx, row in enumerate(t):
                        for c_idx, cell in enumerate(row):
                            tb.cells.append(
                                TableCell(row=r_idx, col=c_idx, text=self._normalize_text(cell or ""))
                            )
                    page_blocks.append(tb.model_dump())

                # 7) Fallback OCR si no hubo texto ni tablas (caso típico: PDF escaneado)
                if self.enable_pdf_ocr_fallback and len(page_blocks) == 0:
                    if pytesseract is None:
                        logger.warning("OCR fallback saltado (pytesseract no disponible) | page=%d", i)
                    else:
                        try:
                            pil_img = page.to_image(resolution=self.ocr_resolution).original
                            data = pytesseract.image_to_data(
                                pil_img, lang=self.ocr_lang, output_type=pytesseract.Output.DICT
                            )
                            words, confs = [], []
                            for w, conf in zip(data.get("text", []), data.get("conf", [])):
                                if not w:
                                    continue
                                words.append(w)
                                try:
                                    confs.append(float(conf))
                                except Exception:
                                    pass
                            ocr_text = self._normalize_text(" ".join(words))
                            mean_conf = (sum(confs) / len(confs) / 100.0) if confs else None

                            # Reconstrucción mínima por líneas OCR (no layout)
                            ocr_lines = [l for l in (ocr_text or "").split("\n") if l.strip()]
                            ocr_paras, ocr_sources = _merge_lines_textual(
                                _normalize_lines_for_merge(ocr_lines, self.dehyphenate, self.normalize_whitespace)
                            )

                            for ptxt, src_idx_list in zip(ocr_paras, ocr_sources):
                                page_blocks.append(
                                    TextBlock(
                                        type="paragraph",
                                        text=ptxt,
                                        ocr=OCRInfo(engine="tesseract", lang=self.ocr_lang,
                                                    dpi=self.ocr_resolution, conf=mean_conf),
                                        prov=Provenance(extractor="pytesseract+merge", stage="ocr-fallback",
                                                        notes=f"source_lines={src_idx_list}"),
                                    ).model_dump()
                                )
                            logger.info("Fallback OCR aplicado | page=%d conf=%.2f", i, (mean_conf or -1))
                        except Exception as e:
                            logger.warning("Fallback OCR falló | page=%d err=%s", i, e)

                # 8) Ensamble de la página IR con métricas por página
                page_ir = PageIR(page_number=i, width=width, height=height, blocks=page_blocks, meta=meta_page)
                pages_ir.append(page_ir)

        return pages_ir

    def _parse_docx(self, path: str) -> List[PageIR]:
        """
        DOCX → PageIR único:
        - Párrafos / Headings detectados por estilo.
        - Tablas con contenido por celda.
        Nota: DOCX no expone páginas físicas; devolvemos una 'página lógica' (page_number=1).
        """
        if docx is None:
            raise ImportError("Instala python-docx para parsear DOCX.")

        document = docx.Document(path)
        page_blocks: List[Dict[str, Any]] = []
        n_raw_paras = 0

        # Párrafos / Headings
        for p in document.paragraphs:
            text = self._normalize_text(p.text or "")
            if not text:
                continue
            n_raw_paras += 1
            style_name = (p.style.name if p.style else "").lower()
            if "heading" in style_name:
                # Detecta nivel (Heading 1, Heading 2, …) de forma simple
                level = 1
                for d in ("1", "2", "3", "4", "5", "6"):
                    if d in style_name:
                        level = int(d)
                        break
                page_blocks.append(
                    TextBlock(type="heading", text=text, level=level,
                              prov=Provenance(extractor="python-docx", stage="parser")).model_dump()
                )
            else:
                if self.enable_list_item_detection and _is_list_item_text(text):
                    cleaned = text.lstrip()
                    cleaned = cleaned[1:].lstrip() if cleaned and cleaned[0] in _BULLETS else cleaned
                    page_blocks.append(
                        TextBlock(type="list_item", text=cleaned,
                                  prov=Provenance(extractor="python-docx", stage="parser")).model_dump()
                    )
                else:
                    page_blocks.append(
                        TextBlock(type="paragraph", text=text,
                                  prov=Provenance(extractor="python-docx", stage="parser")).model_dump()
                    )

        # Tablas
        n_tables = 0
        for tbl in document.tables:
            tb = TableBlock(prov=Provenance(extractor="python-docx", stage="parser"))
            for r_idx, row in enumerate(tbl.rows):
                for c_idx, cell in enumerate(row.cells):
                    tb.cells.append(TableCell(row=r_idx, col=c_idx, text=self._normalize_text(cell.text or "")))
            page_blocks.append(tb.model_dump())
            n_tables += 1

        # Métricas ligeras para DOCX
        meta_page = {
            "n_paragraphs_raw": n_raw_paras,
            "n_blocks_final": len(page_blocks),
            "n_tables": n_tables,
            "layout_loss": 0.0,
        }
        return [PageIR(page_number=1, blocks=page_blocks, meta=meta_page)]

    def _parse_image(self, path: str) -> List[PageIR]:
        """
        IMG → OCR con Tesseract (si está disponible).
        Recomendado: preprocesar con OpenCV (deskew/denoise/binarización) si las imágenes son ruidosas.
        """
        if pytesseract is None or Image is None:
            raise ImportError("Instala pytesseract y Pillow, y Tesseract en el sistema.")

        img = Image.open(path)

        # Igual que en PDF fallback: obtenemos conf promedio
        try:
            data = pytesseract.image_to_data(img, lang=self.ocr_lang, output_type=pytesseract.Output.DICT)
            words, confs = [], []
            for w, conf in zip(data.get("text", []), data.get("conf", [])):
                if not w:
                    continue
                words.append(w)
                try:
                    confs.append(float(conf))
                except Exception:
                    pass
            text = self._normalize_text(" ".join(words))
            mean_conf = (sum(confs) / len(confs) / 100.0) if confs else None
        except Exception:
            # Fallback a image_to_string si falla image_to_data
            raw = pytesseract.image_to_string(img, lang=self.ocr_lang)
            text = self._normalize_text(raw)
            mean_conf = None

        # Reconstrucción mínima por líneas OCR: textual
        lines = [l for l in (text or "").split("\n")]
        norm_lines = _normalize_lines_for_merge(lines, self.dehyphenate, self.normalize_whitespace)
        paras, sources = _merge_lines_textual(norm_lines)
        blocks: List[Dict[str, Any]] = []
        for ptxt, src_idx_list in zip(paras, sources):
            blocks.append(
                TextBlock(
                    type="paragraph",
                    text=ptxt,
                    ocr=OCRInfo(engine="tesseract", lang=self.ocr_lang, dpi=None, conf=mean_conf),
                    prov=Provenance(extractor="pytesseract+merge", stage="parser", notes=f"source_lines={src_idx_list}"),
                ).model_dump()
            )

        meta_page = {
            "n_lines_raw": len([l for l in lines if l.strip()]),
            "n_paragraphs_final": len([b for b in blocks if (isinstance(b, dict) and b.get('type') == 'paragraph')]),
            "layout_loss": 0.0,
        }
        return [PageIR(page_number=1, blocks=blocks, meta=meta_page)]
